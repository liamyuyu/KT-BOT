"""
DOCX Parser
Word 文档解析器
"""

import logging
from pathlib import Path
from .base import BaseParser, ParsedDocument

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """Word 文档解析器"""

    def can_parse(self, file_path: str) -> bool:
        """判断是否为 Word 文件"""
        return file_path.lower().endswith(('.docx', '.doc'))

    async def parse(self, file_path: str) -> ParsedDocument:
        """
        解析 Word 文档

        Args:
            file_path: Word 文件路径

        Returns:
            ParsedDocument: 解析后的文档对象
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        doc = Document(file_path)

        # 提取段落
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)

        # 元数据
        metadata = self._extract_metadata(file_path)
        core_props = doc.core_properties
        metadata.update({
            "author": core_props.author,
            "created": core_props.created.isoformat() if core_props.created else None,
            "modified": core_props.modified.isoformat() if core_props.modified else None
        })

        # 标题提取
        title = core_props.title
        if not title and paragraphs:
            title = paragraphs[0][:100]
        if not title:
            title = Path(file_path).stem

        return ParsedDocument(
            title=title,
            content=content,
            metadata=metadata,
            word_count=len(content.split()),
            file_type="docx"
        )
